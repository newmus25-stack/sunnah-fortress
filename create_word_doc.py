import sys
import subprocess

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    print("Installing python-docx...")
    install("python-docx")
    import docx

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

def add_rtl_paragraph(doc, text, style='Normal', is_heading=False, level=1):
    if is_heading:
        p = doc.add_heading(text, level=level)
    else:
        p = doc.add_paragraph(text, style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Set RTL for all runs
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.rtl = True
        
        # XML level setting for Arabic text logic (complex script)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:cs'), 'Arial')

doc = docx.Document()

# Title
add_rtl_paragraph(doc, "الحديث المتواتر: تعريفه، شروطه، والمقارنة بين الأصوليين والمحدثين", is_heading=True, level=1)

# Introduction
add_rtl_paragraph(doc, "مقدمة:", is_heading=True, level=2)
add_rtl_paragraph(doc, "تستعرض هذه الورقة البحثية مسألة (الحديث المتواتر) من حيث التعريف والشروط، وتوضيح الفرق المنهجي بين علماء أصول الفقه وعلماء الحديث في التعامل مع هذا النوع من الأخبار، وذلك بالاعتماد على مصادر التراث الإسلامي الموثقة.")

add_rtl_paragraph(doc, "أولاً: تعريف المتواتر وشروطه", is_heading=True, level=2)
add_rtl_paragraph(doc, "بالرجوع إلى كتاب (نزهة النظر في توضيح نخبة الفكر) للحافظ ابن حجر العسقلاني (المجلد 1، الصفحة 87):")
add_rtl_paragraph(doc, "- التعريف: هو ما رواه عدد كثير أحالت العادة تواطؤهم أو توافقهم على الكذب، ورووا ذلك عن مثلهم من الابتداء إلى الانتهاء، وكان مستند انتهائهم الحس.")
add_rtl_paragraph(doc, "- الشروط المستخلصة للحكم بالتواتر:")
add_rtl_paragraph(doc, "  1. عدد كثير: لا يتحقق التواتر إلا بجمع من الرواة.")
add_rtl_paragraph(doc, "  2. إحالة العادة التواطؤ على الكذب: أن يمنع العقل والواقع اتفاقهم على اختلاق الخبر.")
add_rtl_paragraph(doc, "  3. الاستمرار في جميع الطبقات: أن يوجد هذا العدد في كل طبقة من طبقات السند.")
add_rtl_paragraph(doc, "  4. أن يكون مستند الخبر الحس: أي أن يقول الرواة (سمعنا) أو (رأينا).")

add_rtl_paragraph(doc, "ثانياً: المتواتر عند المحدثين", is_heading=True, level=2)
add_rtl_paragraph(doc, "بالرجوع إلى كتاب (تدريب الراوي في شرح تقريب النواوي) للإمام جلال الدين السيوطي (المجلد 2، الصفحة 621):")
add_rtl_paragraph(doc, "- النص: يذكر الإمام السيوطي في النوع الثلاثين (المشهور): «ومنه المتواتر المعروف في الفقه وأصوله، ولا يذكره المحدثون، وهو قليل لا يكاد يوجد في رواياتهم، وهو ما نقله من يحصل العلم بصدقهم ضرورة عن مثلهم من أوله إلى آخره».")
add_rtl_paragraph(doc, "- التوضيح: يبيّن الإمام السيوطي أن المتواتر ليس من المباحث الأساسية عند علماء الحديث (بالمعنى الاصطلاحي الدقيق لعلم الإسناد)، والسبب في ذلك أن المتواتر يفيد (العلم الضروري) اليقيني الذي لا يحتاج معه الباحث إلى تفتيش أحوال الرواة أو الجرح والتعديل.")

add_rtl_paragraph(doc, "ثالثاً: المقارنة بين الأصوليين والمحدثين", is_heading=True, level=2)
add_rtl_paragraph(doc, "من خلال تتبع المادة في المصدرين السابقين ومقارنة المناهج، يمكن تلخيص الفروق فيما يلي:")
add_rtl_paragraph(doc, "1. زاوية النظر: الأصوليون ينظرون إليه كدليل يفيد العلم القطعي واليقين لإثبات الأحكام والعقائد. بينما المحدثون ينظرون إليه من حيث السند وطرق الرواية وتمييزه عن أخبار الآحاد.")
add_rtl_paragraph(doc, "2. مكانة المصطلح: هو من صلب مباحث الأخبار في كتب الأصول، بينما يعتبره بعض المحدثين خارجاً عن (علم الإسناد والرجال) لأنه لا يحتاج لبحث رجاله لقطعيته.")
add_rtl_paragraph(doc, "3. الوجود الواقعي: الأصوليون يتوسعون في شروطه المنطقية والعقلية. بينما المحدثون (كابن الصلاح) رأوا أنه عزيز الوجود جداً في السنة، وقد رد عليه ابن حجر والسيوطي بأنه موجود بكثرة وصنف السيوطي في ذلك (قطف الأزهار المتناثرة في الأخبار المتواترة).")

add_rtl_paragraph(doc, "الخلاصة:", is_heading=True, level=2)
add_rtl_paragraph(doc, "المتواتر عند الأصوليين هو الخبر الذي يقطع بصدقه، بينما المحدثون وإن وافقوا على هذا التعريف إلا أنهم لم يوسعوا القول فيه في كتبهم الأولى لأن غرضهم هو نقد الأسانيد، والمتواتر فوق مستوى النقد لقطعيته.")

doc.save(r'c:\antigravivty\الحديث_المتواتر.docx')
print("Document saved successfully at c:\\antigravivty\\الحديث_المتواتر.docx")
